#!/usr/bin/env python3
"""
test_apis.py — Validate Legistar API + Pinecone connectivity before running the full job.

Run:
    pip install -r requirements.txt
    cp env.example .env   # fill in your real keys
    python test_apis.py
"""

import os
import sys
import json
import requests
from dotenv import load_dotenv

load_dotenv()

# ── Configuration ────────────────────────────────────────────────────────────
LEGISTAR_BASE = "https://webapi.legistar.com/v1"
CLIENT = "pittsburgh"
CLIENT_LABEL = "Pittsburgh City Council"

PINECONE_API_KEY = os.environ.get("PINECONE_API_KEY", "")
PINECONE_INDEX = os.environ.get("PINECONE_INDEX_NAME", "")
PINECONE_NAMESPACE = os.environ.get("PINECONE_NAMESPACE", "legislation")

passed = 0
failed = 0


def banner(msg: str):
    print(f"\n{'='*60}\n  {msg}\n{'='*60}")


def ok(msg: str):
    global passed
    passed += 1
    print(f"  ✅  {msg}")


def fail(msg: str):
    global failed
    failed += 1
    print(f"  ❌  {msg}")


# ── Test 1: Legistar API — list matters ──────────────────────────────────────
def test_legistar_matters():
    banner("Test 1 · Legistar REST API — fetch matters")

    url = f"{LEGISTAR_BASE}/{CLIENT}/matters?$top=3&$orderby=MatterIntroDate desc"
    print(f"\n  → GET {url}")
    try:
        resp = requests.get(url, timeout=15)
        resp.raise_for_status()
        data = resp.json()
        if isinstance(data, list) and len(data) > 0:
            m = data[0]
            ok(f"{CLIENT_LABEL}: got {len(data)} matters. "
               f"Latest: [{m.get('MatterFile')}] {(m.get('MatterTitle') or '')[:80]}")
        else:
            fail("Response was empty or unexpected format")
    except Exception as exc:
        fail(f"Legistar matters endpoint failed: {exc}")


# ── Test 2: Legistar API — date filtering ────────────────────────────────────
def test_legistar_date_filter():
    banner("Test 2 · Legistar date-range filtering (2022–2025)")

    url = (
        f"{LEGISTAR_BASE}/{CLIENT}/matters"
        f"?$filter=MatterIntroDate ge datetime'2022-01-01'"
        f" and MatterIntroDate lt datetime'2026-01-01'"
        f"&$top=5&$orderby=MatterIntroDate asc"
    )
    print(f"\n  → GET {url}")
    try:
        resp = requests.get(url, timeout=15)
        resp.raise_for_status()
        data = resp.json()
        if isinstance(data, list) and len(data) > 0:
            dates = [m.get("MatterIntroDate", "?")[:10] for m in data]
            ok(f"Date filter works. First dates returned: {dates}")
        else:
            fail("Date filter returned no results")
    except Exception as exc:
        fail(f"Date filter test failed: {exc}")


# ── Test 3: Legistar API — MatterTitle content quality ───────────────────────
def test_legistar_title_quality():
    banner("Test 3 · MatterTitle text content quality")

    url = (
        f"{LEGISTAR_BASE}/{CLIENT}/matters"
        f"?$filter=MatterIntroDate ge datetime'2024-01-01'"
        f" and MatterTypeName eq 'Resolution'"
        f"&$top=5&$orderby=MatterIntroDate desc"
    )
    try:
        resp = requests.get(url, timeout=15)
        resp.raise_for_status()
        data = resp.json()
        lengths = [len(m.get("MatterTitle") or "") for m in data]
        avg_len = sum(lengths) / len(lengths) if lengths else 0
        ok(f"Avg MatterTitle length: {avg_len:.0f} chars "
           f"(range: {min(lengths)}-{max(lengths)})")
        print(f"        Sample: {(data[0].get('MatterTitle') or '')[:120]}...")
    except Exception as exc:
        fail(f"Title quality test failed: {exc}")


# ── Test 4: Legistar API — attachments ───────────────────────────────────────
def test_legistar_attachments():
    banner("Test 4 · Legistar attachments endpoint")

    url = (
        f"{LEGISTAR_BASE}/{CLIENT}/matters"
        f"?$filter=MatterIntroDate ge datetime'2024-06-01'"
        f"&$top=10&$orderby=MatterIntroDate desc"
    )
    try:
        resp = requests.get(url, timeout=15)
        resp.raise_for_status()
        data = resp.json()

        with_attach = 0
        sample_url = None
        for m in data:
            aurl = f"{LEGISTAR_BASE}/{CLIENT}/matters/{m['MatterId']}/attachments"
            aresp = requests.get(aurl, timeout=10)
            raw = aresp.json()
            if isinstance(raw, list) and raw:
                with_attach += 1
                if not sample_url:
                    sample_url = raw[0].get("MatterAttachmentHyperlink", "")
        ok(f"{with_attach}/{len(data)} matters have attachments")
        if sample_url:
            print(f"        Sample attachment URL: {sample_url}")
    except Exception as exc:
        fail(f"Attachments test failed: {exc}")


# ── Test 5: Attachment download ──────────────────────────────────────────────
def test_attachment_download():
    banner("Test 5 · Download and parse a sample attachment")

    url = (
        f"{LEGISTAR_BASE}/{CLIENT}/matters"
        f"?$filter=MatterIntroDate ge datetime'2024-06-01'"
        f"&$top=20&$orderby=MatterIntroDate desc"
    )
    try:
        resp = requests.get(url, timeout=15)
        resp.raise_for_status()
        data = resp.json()

        # Find a PDF attachment
        for m in data:
            aurl = f"{LEGISTAR_BASE}/{CLIENT}/matters/{m['MatterId']}/attachments"
            aresp = requests.get(aurl, timeout=10)
            raw = aresp.json()
            if not isinstance(raw, list):
                continue
            for att in raw:
                link = att.get("MatterAttachmentHyperlink", "")
                name = att.get("MatterAttachmentName", "")
                if link.lower().endswith(".pdf"):
                    print(f"  → Downloading: {name}")
                    print(f"    URL: {link}")
                    dresp = requests.get(link, timeout=30)
                    if dresp.status_code == 200 and len(dresp.content) > 100:
                        from pypdf import PdfReader
                        import io
                        reader = PdfReader(io.BytesIO(dresp.content))
                        text = ""
                        for page in reader.pages[:3]:  # first 3 pages
                            text += page.extract_text() or ""
                        if text.strip():
                            ok(f"PDF parsed: {len(reader.pages)} pages, "
                               f"extracted {len(text)} chars from first 3")
                            print(f"        Preview: {text[:150]}...")
                        else:
                            ok(f"PDF downloaded ({len(dresp.content)} bytes) "
                               f"but no extractable text (scanned image PDF)")
                        return
                    else:
                        print(f"  ⚠️  Download returned status {dresp.status_code}")

                elif link.lower().endswith(".docx"):
                    print(f"  → Downloading: {name}")
                    dresp = requests.get(link, timeout=30)
                    if dresp.status_code == 200:
                        from docx import Document
                        import io
                        doc = Document(io.BytesIO(dresp.content))
                        text = "\n".join(p.text for p in doc.paragraphs[:20])
                        if text.strip():
                            ok(f"DOCX parsed: {len(doc.paragraphs)} paragraphs, "
                               f"extracted {len(text)} chars")
                            print(f"        Preview: {text[:150]}...")
                        else:
                            ok("DOCX downloaded but had no text paragraphs")
                        return

        fail("Could not find a downloadable PDF or DOCX attachment in sample")
    except Exception as exc:
        fail(f"Attachment download test failed: {exc}")


# ── Test 6: Pinecone connection ──────────────────────────────────────────────
def test_pinecone_connection():
    banner("Test 6 · Pinecone connection")

    if not PINECONE_API_KEY:
        fail("PINECONE_API_KEY not set in environment / .env")
        return
    if not PINECONE_INDEX:
        fail("PINECONE_INDEX_NAME not set in environment / .env")
        return

    try:
        from pinecone import Pinecone
        pc = Pinecone(api_key=PINECONE_API_KEY)
        idx = pc.Index(PINECONE_INDEX)
        stats = idx.describe_index_stats()
        ok(f"Connected to index '{PINECONE_INDEX}'. "
           f"Namespaces: {list(stats.get('namespaces', {}).keys()) or ['(empty)']}, "
           f"Total vectors: {stats.get('total_vector_count', 0)}")
    except Exception as exc:
        fail(f"Pinecone connection failed: {exc}")


# ── Test 7: Pinecone test upsert + delete ────────────────────────────────────
def test_pinecone_upsert():
    banner("Test 7 · Pinecone test upsert (integrated inference)")

    if not PINECONE_API_KEY or not PINECONE_INDEX:
        fail("Skipped — Pinecone env vars not set")
        return

    try:
        from pinecone import Pinecone
        pc = Pinecone(api_key=PINECONE_API_KEY)
        idx = pc.Index(PINECONE_INDEX)

        test_record = {
            "_id": "__test_record__",
            "text": (
                "Resolution authorizing the Mayor and Director of Finance "
                "to enter into an agreement for auditing services for the "
                "City of Pittsburgh for the fiscal year 2025."
            ),
            "title": "Test Record — Safe to Delete",
            "type": "legislation",
            "date": "2025-01-01",
            "url": "https://example.com/test",
            "source": "Test Script",
            "tags": ["test"],
            "summary": "Temporary test record for API validation.",
        }

        print("  → Upserting test record via upsert_records()...")
        result = idx.upsert_records(
            namespace=PINECONE_NAMESPACE,
            records=[test_record],
        )
        ok(f"upsert_records() succeeded: {result}")

        # Clean up
        print("  → Deleting test record...")
        idx.delete(ids=["__test_record__"], namespace=PINECONE_NAMESPACE)
        ok("Test record deleted")

    except AttributeError:
        fail("upsert_records() not available — ensure pinecone>=5.0.0 "
             "and your index is configured for integrated inference")
    except Exception as exc:
        fail(f"Pinecone upsert test failed: {exc}")


# ── Summary ──────────────────────────────────────────────────────────────────
def main():
    print("\n🔍  Running API validation tests...\n")

    test_legistar_matters()
    test_legistar_date_filter()
    test_legistar_title_quality()
    test_legistar_attachments()
    test_attachment_download()
    test_pinecone_connection()
    test_pinecone_upsert()

    banner("RESULTS")
    print(f"  ✅  Passed: {passed}")
    print(f"  ❌  Failed: {failed}")
    if failed:
        print("\n  ⚠️  Fix the failures above before running the full scraping job.")
        sys.exit(1)
    else:
        print("\n  🎉  All tests passed! You're ready to run scrape_legislation.py")
        sys.exit(0)


if __name__ == "__main__":
    main()
